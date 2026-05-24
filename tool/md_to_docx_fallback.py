import re
import sys
from pathlib import Path
from docx import Document


def convert(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()
    doc = Document()
    para = []
    in_code = False

    def flush_para():
        nonlocal para
        if para:
            doc.add_paragraph(' '.join(para).strip())
            para = []

    for line in lines:
        s = line.rstrip('\n')

        if s.strip().startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            continue

        if not s.strip():
            flush_para()
            continue

        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m:
            flush_para()
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            continue

        m = re.match(r'^\s*[-*+]\s+(.*)$', s)
        if m:
            flush_para()
            doc.add_paragraph(m.group(1).strip(), style='List Bullet')
            continue

        m = re.match(r'^\s*\d+[\.)]\s+(.*)$', s)
        if m:
            flush_para()
            doc.add_paragraph(m.group(1).strip(), style='List Number')
            continue

        para.append(s.strip())

    flush_para()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == '__main__':
    args = sys.argv[1:]
    if len(args) % 2 != 0 or not args:
        raise SystemExit('Use: <src1> <dst1> ...')

    for i in range(0, len(args), 2):
        src = Path(args[i])
        dst = Path(args[i + 1])
        convert(src, dst)
        print(f'OK {dst.as_posix()}')
